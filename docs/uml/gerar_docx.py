# -*- coding: utf-8 -*-
"""
Gerador de documentos Word (.docx) em formato ABNT para o TCC ECOnecta.
Converte o conteudo Markdown existente + embute as imagens dos diagramas.

ABNT aplicado: A4, margens 3/2/3/2 cm, Times New Roman 12, espacamento 1,5,
paragrafos justificados com recuo de primeira linha 1,25 cm, titulos numerados,
figuras com legenda (superior) e fonte (inferior), tabelas com grade.
"""
import os
import re
import struct
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))          # docs/uml
DOCS = os.path.dirname(BASE)                                 # docs
PNG = os.path.join(BASE, "imagens-png")
OUT = os.path.join(BASE, "word")
os.makedirs(OUT, exist_ok=True)

FONT = "Times New Roman"
CONTENT_W_CM = 15.5
MAX_H_CM = 20.5

# ---------------------------------------------------------------- helpers ----

def png_size(path):
    with open(path, "rb") as f:
        head = f.read(24)
    w, h = struct.unpack(">II", head[16:24])
    return w, h

def setup_abnt(doc):
    # A4 + margens ABNT
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(3.0)
        sec.left_margin = Cm(3.0)
        sec.bottom_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
    # estilo Normal
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # estilos de titulo
    for i, sz in [(1, 12), (2, 12), (3, 12), (4, 12)]:
        st = doc.styles["Heading %d" % i]
        st.font.name = FONT
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        st.paragraph_format.keep_with_next = True

def body_par(doc, text="", indent=True, justify=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if text:
        add_inline(p, text)
    return p

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

def clean_inline(text):
    text = LINK_RE.sub(r"\1", text)          # [txt](url) -> txt
    text = text.replace("\\|", "|")
    return text

def add_inline(par, text):
    text = clean_inline(text)
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            par.add_run(part)

# contadores globais de figura/tabela
counters = {"fig": 0, "tab": 0}

def add_figure(doc, img_path, titulo, descricao=None, fonte="Elaborado pelo autor (2026)."):
    counters["fig"] += 1
    n = counters["fig"]
    cap = doc.add_paragraph()
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    r = cap.add_run("Figura %d – %s" % (n, titulo))
    r.font.size = Pt(10); r.font.name = FONT
    # imagem
    try:
        w, h = png_size(img_path)
        ratio = h / w
        dw = CONTENT_W_CM
        dh = dw * ratio
        if dh > MAX_H_CM:
            dh = MAX_H_CM; dw = dh / ratio
        ip = doc.add_paragraph()
        ip.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(img_path, width=Cm(dw))
    except Exception as e:
        body_par(doc, "[imagem nao encontrada: %s]" % os.path.basename(img_path))
    src = doc.add_paragraph()
    src.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(6)
    r = src.add_run("Fonte: %s" % fonte)
    r.font.size = Pt(10); r.font.name = FONT
    if descricao:
        body_par(doc, descricao)

# ------------------------------------------------------ conversor markdown ----

def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ""
            txt = row[ci] if ci < len(row) else ""
            par = cell.paragraphs[0]
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(par, txt)
            for run in par.runs:
                run.font.size = Pt(10)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def is_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$", line))

def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # nao quebrar em \| escapado
    cells = re.split(r"(?<!\\)\|", line)
    return [c.strip() for c in cells]

def convert_md(doc, md, skip_h1=False, base_heading=0):
    lines = md.split("\n")
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        # blocos de codigo: pular inteiro
        if line.lstrip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue
        # imagens
        m = IMG_RE.match(line.strip())
        if m:
            alt, path = m.group(1), m.group(2)
            ip = os.path.join(os.path.dirname(md_path_ctx[0]), path)
            add_figure(doc, ip, alt or "Diagrama")
            i += 1
            continue
        # tabela
        if line.strip().startswith("|") and i + 1 < len(lines) and is_sep(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue
        # titulos
        hm = re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1)) + base_heading
            txt = clean_inline(hm.group(2)).strip()
            if level == 1 and skip_h1:
                i += 1
                continue
            level = min(level, 4)
            h = doc.add_heading(level=level)
            add_inline(h, txt)
            i += 1
            continue
        # regua / vazio
        if re.match(r"^\s*([-*_])\1{2,}\s*$", line):
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        # blockquote
        if line.lstrip().startswith(">"):
            p = body_par(doc, line.lstrip()[1:].strip(), indent=False)
            for r in p.runs:
                r.italic = True
            i += 1
            continue
        # lista nao ordenada
        bm = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline(p, bm.group(1))
            i += 1
            continue
        # lista ordenada
        om = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if om:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline(p, om.group(1))
            i += 1
            continue
        # paragrafo: agrupa linhas ate vazio
        buf = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6})\s", lines[i]) \
                and not lines[i].strip().startswith("|") and not lines[i].lstrip().startswith(("-", "*", ">", "```")) \
                and not re.match(r"^\s*\d+\.\s", lines[i]):
            buf.append(lines[i].strip())
            i += 1
        body_par(doc, " ".join(buf))

md_path_ctx = [BASE]

def read(p):
    with open(p, "r", encoding="utf-8") as f:
        return f.read()

def add_md_file(doc, path, skip_h1=False, base_heading=0):
    md_path_ctx[0] = path
    convert_md(doc, read(path), skip_h1=skip_h1, base_heading=base_heading)

def title_page_heading(doc, titulo, subtitulo=None):
    h = doc.add_heading(level=1)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(h, titulo)
    if subtitulo:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitulo); r.italic = True
    doc.add_paragraph()

# ----------------------------------------------------- catalogo de figuras ----
# (titulo, arquivo png, descricao curta)
FIGS = {
 "Casos de Uso": [
   ("1_casos-de-uso.png", "Diagrama de Casos de Uso do ECOnecta com os atores Cidadao, Empresa Coletora, Administrador e os sistemas externos Cloudinary e ViaCEP."),
 ],
 "Classes de Dados": [
   ("2_classes-de-dados.png", "Diagrama de Classes de Dados com as 12 entidades do dominio e suas multiplicidades, derivado do schema Prisma."),
 ],
 "Diagramas de Atividades": [
   ("3_atividade-macro.png", "Atividade do fluxo macro: criacao da solicitacao, moderacao, aceite, execucao e avaliacao da coleta."),
   ("3_atividade-auth.png", "Atividade do fluxo de autenticacao web e redirecionamento por perfil."),
   ("3_atividade-cancelar.png", "Atividade do cancelamento de solicitacao e suas regras de estado."),
 ],
 "Diagramas de Sequencia": [
   ("4_SEQ-01-login-web.png", "Sequencia da autenticacao web (NextAuth, JWT)."),
   ("4_SEQ-02-login-mobile.png", "Sequencia da autenticacao mobile (login e refresh de tokens)."),
   ("4_SEQ-03-cadastro.png", "Sequencia do cadastro de usuario/empresa."),
   ("4_SEQ-04-recuperar-senha.png", "Sequencia da recuperacao de senha (forgot e reset)."),
   ("4_SEQ-05-criar-solicitacao.png", "Sequencia da criacao de solicitacao com upload de imagens."),
   ("4_SEQ-06-listar-solicitacoes.png", "Sequencia da listagem de solicitacoes com ramificacao por perfil."),
   ("4_SEQ-07-moderar-solicitacao.png", "Sequencia da moderacao (aprovar/rejeitar) pelo administrador."),
   ("4_SEQ-08-cancelar-solicitacao.png", "Sequencia do cancelamento de solicitacao."),
   ("4_SEQ-09-negociar-pre-aceite.png", "Sequencia da negociacao (chat pre-aceite) por solicitacao."),
   ("4_SEQ-10-aceitar-solicitacao.png", "Sequencia do aceite da solicitacao com criacao da coleta em transacao."),
   ("4_SEQ-11-atualizar-coleta.png", "Sequencia da atualizacao/conclusao do status da coleta."),
   ("4_SEQ-12-chat-coleta.png", "Sequencia do chat operacional vinculado a coleta."),
   ("4_SEQ-13-avaliar-coleta.png", "Sequencia da avaliacao da coleta concluida."),
   ("4_SEQ-14-notificacoes-sse.png", "Sequencia das notificacoes em tempo real via Server-Sent Events."),
   ("4_SEQ-15-consultar-cep.png", "Sequencia da consulta de CEP (integracao ViaCEP)."),
 ],
 "Diagrama de Entidade e Relacionamento": [
   ("5_er-econecta.png", "Modelo Entidade-Relacionamento com as 12 tabelas fisicas em PostgreSQL."),
 ],
 "Classes Participantes (Robustez)": [
   ("7_CP-01-cadastro.png", "Robustez do cadastro de usuario/empresa."),
   ("7_CP-02-autenticar.png", "Robustez da autenticacao."),
   ("7_CP-03-criar-solicitacao.png", "Robustez da criacao de solicitacao."),
   ("7_CP-04-moderar-solicitacao.png", "Robustez da moderacao de solicitacao."),
   ("7_CP-05-negociar-pre-aceite.png", "Robustez da negociacao pre-aceite."),
   ("7_CP-06-aceitar-solicitacao.png", "Robustez do aceite da solicitacao."),
   ("7_CP-07-atualizar-coleta.png", "Robustez da atualizacao/conclusao da coleta."),
   ("7_CP-08-chat-coleta.png", "Robustez do chat da coleta."),
   ("7_CP-09-avaliar-coleta.png", "Robustez da avaliacao da coleta."),
   ("7_CP-10-notificacoes-sse.png", "Robustez das notificacoes (SSE)."),
 ],
 "Diagramas de Arquitetura": [
   ("8_componentes-econecta.png", "Diagrama de Componentes da aplicacao."),
   ("8_pacotes-econecta.png", "Diagrama de Pacotes / organizacao do codigo."),
   ("8_deploy-econecta.png", "Diagrama de Implantacao (deploy inferido)."),
   ("8_classes-econecta.png", "Diagrama de Classes com a camada de servicos."),
   ("8_estado-solicitacao.png", "Maquina de estados da Solicitacao."),
   ("8_estado-coleta.png", "Maquina de estados da Coleta."),
   ("8_estado-conversa.png", "Maquina de estados da Conversa de pre-aceite."),
 ],
}

def add_appendix_figures(doc, heading_prefix="APENDICE"):
    letras = "ABCDEFGHIJ"
    for idx, (cat, items) in enumerate(FIGS.items()):
        doc.add_heading("%s %s – %s" % (heading_prefix, letras[idx], cat), level=1)
        for fname, desc in items:
            if len(items) == 1:
                titulo = cat
            else:
                titulo = fname.split("_", 1)[1].replace(".png", "").replace("-", " ").strip()
            add_figure(doc, os.path.join(PNG, fname), titulo, desc)

# =========================================================== DOCUMENTOS =======

def doc1_completa():
    d = Document(); setup_abnt(d)
    title_page_heading(d, "Documentação UML — ECOnecta",
                       "Plataforma de solicitação e coleta de materiais recicláveis")
    add_md_file(d, os.path.join(BASE, "documentacao-uml-completa.md"), skip_h1=True)
    d.add_page_break()
    d.add_heading("Apêndices — Diagramas", level=1)
    add_appendix_figures(d)
    out = os.path.join(OUT, "1-ECOnecta-Documentacao-UML-completa-ABNT.docx")
    d.save(out); return out

def doc2_apendices():
    d = Document(); setup_abnt(d)
    title_page_heading(d, "Apêndices de Diagramas — ECOnecta")
    add_appendix_figures(d)
    out = os.path.join(OUT, "2-ECOnecta-Apendices-Diagramas-ABNT.docx")
    d.save(out); return out

def doc3_dicionario():
    d = Document(); setup_abnt(d)
    add_md_file(d, os.path.join(BASE, "6-dicionario-de-dados", "dicionario-de-dados.md"))
    out = os.path.join(OUT, "3-ECOnecta-Dicionario-de-Dados-ABNT.docx")
    d.save(out); return out

def doc4_casos_uso():
    d = Document(); setup_abnt(d)
    title_page_heading(d, "Descrições de Casos de Uso — ECOnecta")
    add_md_file(d, os.path.join(DOCS, "APENDICE-B-descricoes-casos-de-uso.md"), skip_h1=True)
    d.add_heading("Casos de uso adicionais (UC22–UC25 e UC15)", level=1)
    add_md_file(d, os.path.join(BASE, "1-casos-de-uso", "descricoes-casos-de-uso.md"),
                skip_h1=True, base_heading=0)
    out = os.path.join(OUT, "4-ECOnecta-Descricoes-Casos-de-Uso-ABNT.docx")
    d.save(out); return out

if __name__ == "__main__":
    for fn in (doc1_completa, doc2_apendices, doc3_dicionario, doc4_casos_uso):
        counters["fig"] = 0; counters["tab"] = 0
        path = fn()
        print("OK:", os.path.relpath(path, BASE), "(%.0f KB)" % (os.path.getsize(path) / 1024))
